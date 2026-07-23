# -*- coding: utf-8 -*-
"""Generate GC-KB-0001 Account Creation (stylized .docx). Reproducible KB builder."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "..", "..", "GC-KB-0001-accounts-account-creation-v1.0.docx")

# ---- palette (warm gold / bronze — matches the app's default "marble" theme) ----
INK       = RGBColor(0x2B, 0x26, 0x20)   # warm near-black body text
MUTED     = RGBColor(0x6B, 0x63, 0x55)   # muted captions
BRONZE     = RGBColor(0x5C, 0x4A, 0x1E)   # deep bronze headings
GOLD      = RGBColor(0x8A, 0x6D, 0x1F)   # gold accent
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
TEAL      = RGBColor(0x0F, 0x76, 0x6E)   # "new user" perspective
SLATE     = RGBColor(0x47, 0x55, 0x69)   # "system" perspective
GREEN     = RGBColor(0x15, 0x80, 0x3D)
AMBERTXT  = RGBColor(0x7A, 0x5B, 0x14)

H_BAND    = "5C4A1E"   # heading band / table header fill
H_TITLE   = "463914"   # title band fill (darker)
ROW_ALT   = "FAF6EC"   # light row shade
CALL_AMBER= "FDF6E3"   # amber callout bg
CALL_GREEN= "EAF4EC"   # green callout bg
CALL_BLUE = "EEF2F7"   # neutral note bg
BORDER    = "D9CDA8"   # light table border

FIG = {"n": 0}

def _shade(el, fill):
    tcPr = el._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

def write_cell(cell, text, bold=False, color=INK, size=10, align=None, italic=False, font='Calibri'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    parts = text.split('\n')
    for i, seg in enumerate(parts):
        if i:
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
            if align is not None:
                p.alignment = align
        r = p.add_run(seg)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = font
    return cell

def table_borders(tbl, color=BORDER, sz=4):
    t = tbl._tbl
    tblPr = t.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def para(doc, text='', size=10.5, color=INK, bold=False, italic=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    if align is not None: p.alignment = align
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold; r.italic = italic; r.font.name = 'Calibri'
    return p

def rich(doc, segments, size=10.5, space_after=6, space_before=0):
    """segments: list of (text, bold, color)."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    for text, bold, color in segments:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color; r.font.name = 'Calibri'
    return p

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BRONZE; r.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12'); bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'C99A2E')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(doc, text, chip=None, chip_color=GOLD):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    if chip:
        cr = p.add_run(chip + '  '); cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = chip_color; cr.font.name = 'Calibri'
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK; r.font.name = 'Calibri'
    return p

def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = INK; r.font.name = 'Calibri'
            r2 = p.add_run(rest); r2.font.size = Pt(size); r2.font.color.rgb = INK; r2.font.name = 'Calibri'
        else:
            r = p.add_run(it); r.font.size = Pt(size); r.font.color.rgb = INK; r.font.name = 'Calibri'

def callout(doc, title, body, bg, bar_hex, title_color):
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.cell(0, 0); _shade(cell, bg); _set_cell_margins(cell, 90, 90, 140, 140)
    # left accent bar
    tcPr = cell._tc.get_or_add_tcPr(); borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left'); left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24'); left.set(qn('w:space'), '0'); left.set(qn('w:color'), bar_hex)
    borders.append(left); tcPr.append(borders)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = title_color; r.font.name = 'Calibri'
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body); r2.font.size = Pt(10); r2.font.color.rgb = INK; r2.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def step_table(doc, steps, perspective, persp_color_hex):
    tbl = doc.add_table(rows=1, cols=2); table_borders(tbl)
    hdr = tbl.rows[0].cells
    _shade(hdr[0], persp_color_hex); _shade(hdr[1], persp_color_hex)
    write_cell(hdr[0], '#', bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(hdr[1], perspective, bold=True, color=WHITE, size=10)
    for i, s in enumerate(steps):
        row = tbl.add_row().cells
        write_cell(row[0], str(s[0]), bold=True, color=BRONZE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row[1], s[1], size=10.5)
        if i % 2 == 1:
            _shade(row[0], ROW_ALT); _shade(row[1], ROW_ALT)
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tbl.columns[0].width = Inches(0.4)
    tbl.columns[1].width = Inches(6.4)
    for r in tbl.rows:
        r.cells[0].width = Inches(0.4); r.cells[1].width = Inches(6.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def figure(doc, filename, caption):
    FIG["n"] += 1
    path = os.path.join(HERE, filename)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(); run.add_picture(path, width=Inches(6.1))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after = Pt(10)
    r = cp.add_run("Figure %d — %s" % (FIG["n"], caption)); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED; r.font.name = 'Calibri'

def kv_table(doc, rows, label_w=1.9, col_hdr=None):
    tbl = doc.add_table(rows=0, cols=2); table_borders(tbl)
    if col_hdr:
        hc = tbl.add_row().cells; _shade(hc[0], H_BAND); _shade(hc[1], H_BAND)
        write_cell(hc[0], col_hdr[0], bold=True, color=WHITE, size=10)
        write_cell(hc[1], col_hdr[1], bold=True, color=WHITE, size=10)
    for i, (k, v) in enumerate(rows):
        rc = tbl.add_row().cells
        write_cell(rc[0], k, bold=True, color=BRONZE, size=10)
        write_cell(rc[1], v, size=10)
        if i % 2 == 1:
            _shade(rc[0], ROW_ALT); _shade(rc[1], ROW_ALT)
    for r in tbl.rows:
        r.cells[0].width = Inches(label_w); r.cells[1].width = Inches(6.8 - label_w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def three_col(doc, header, rows, widths):
    tbl = doc.add_table(rows=0, cols=3); table_borders(tbl)
    hc = tbl.add_row().cells
    for j in range(3):
        _shade(hc[j], H_BAND); write_cell(hc[j], header[j], bold=True, color=WHITE, size=10)
    for i, row in enumerate(rows):
        rc = tbl.add_row().cells
        for j in range(3):
            write_cell(rc[j], row[j], size=10, bold=(j == 0), color=(BRONZE if j == 0 else INK))
        if i % 2 == 1:
            for j in range(3): _shade(rc[j], ROW_ALT)
    for r in tbl.rows:
        for j in range(3): r.cells[j].width = Inches(widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def add_page_number(paragraph):
    def fld(instr):
        r = OxmlElement('w:r'); f = OxmlElement('w:fldSimple'); f.set(qn('w:instr'), instr); r.append(f); return f
    run = paragraph.add_run(); run.font.size = Pt(8); run.font.color.rgb = MUTED
    fldSimple = OxmlElement('w:fldSimple'); fldSimple.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fldSimple)

# ============================ BUILD ============================
doc = Document()
# base style
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5); st.font.color.rgb = INK
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

# core properties
cp = doc.core_properties
cp.title = "GC-KB-0001 Account Creation"
cp.author = "Gospel Central Operations"
cp.category = "accounts"
cp.comments = "Operations Knowledge Base article. Verified against build 5ce6222."

# ---- footer ----
footer = sec.footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("GC-KB-0001  ·  Account Creation  ·  v1.0  ·  Gospel Central Operations KB  ·  Page ")
fr.font.size = Pt(8); fr.font.color.rgb = MUTED; fr.font.name = 'Calibri'
add_page_number(fp)

# ---- TITLE BAND ----
band = doc.add_table(rows=1, cols=1); cell = band.cell(0, 0); _shade(cell, H_TITLE); _set_cell_margins(cell, 160, 160, 200, 200)
cell.text = ''
p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
r = p.add_run("GOSPEL CENTRAL  ·  OPERATIONS KNOWLEDGE BASE"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xE7,0xD4,0x9A); r.font.name = 'Calibri'
p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2); p2.paragraph_format.space_before = Pt(2)
r2 = p2.add_run("Account Creation"); r2.bold = True; r2.font.size = Pt(26); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'
p3 = cell.add_paragraph(); p3.paragraph_format.space_after = Pt(0)
r3 = p3.add_run("How an account goes from the creator to the newly-created user — every step, in plain English."); r3.italic = True; r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0xEC,0xE2,0xC8); r3.font.name = 'Calibri'
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ---- METADATA ----
meta = [
    ("Article ID", "GC-KB-0001"),
    ("Title", "Account Creation"),
    ("Category", "accounts"),
    ("Version", "v1.0"),
    ("Status", "Draft (pending review)"),
    ("Owner", "Gospel Central Operations"),
    ("Created", "2026-07-22"),
    ("Last updated", "2026-07-22"),
    ("Applies-to build", "5ce6222 (verified against source at this commit)"),
    ("Source of truth", "Codebase — see the “Source of Truth” section at the end"),
]
kv_table(doc, meta, label_w=1.7)

# ---- AT A GLANCE ----
callout(doc, "At a glance",
        "An authorized leader opens Admin → Users → Add User and fills a 3-step wizard (identity → role & “reports to” → review). "
        "On “Create account” the system checks permissions, creates the user, and generates a one-time temporary password shown to the creator exactly once. "
        "The creator shares the username + temporary password with the new person, who signs in with it, is forced to set their own password, and then has full access. "
        "The temporary password stops working once the new password is set.",
        CALL_BLUE, "5C4A1E", BRONZE)

# ---- WHO CAN CREATE ----
h1(doc, "1. Who can create accounts")
para(doc, "Creating a user is a leader-tier action. Members can never create accounts. A creator can only assign roles below their own level, and only Branch Leaders and above may create people outside their own group (“cross-branch caretaking”).")
three_col(doc, ["Creator role", "Can create", "Where (scope)"],
    [
        ["Developer", "Any role, including another Developer", "Anywhere"],
        ["Overseer", "Any role below Overseer", "Anywhere"],
        ["Branch Leader", "Any role below Branch Leader", "Any branch (cross-branch caretaking)"],
        ["Group Leader", "Any role below Group Leader", "Only within their own group (subtree)"],
        ["Team Leader", "Any role below Team Leader", "Only within their own team (subtree)"],
        ["Member", "— cannot create accounts", "—"],
    ], widths=[1.5, 2.7, 2.6])
callout(doc, "Nuance worth knowing",
        "The “Add User” button appears for Team Leader and above using a role-only check, but the real scope limit (own group vs. anywhere) is enforced by the permission gate and the server when the account is submitted. So a Team or Group Leader sees the button, yet can still only create people inside their own group.",
        CALL_AMBER, "C99A2E", AMBERTXT)

# ---- THE PROCESS ----
h1(doc, "2. The end-to-end process")

# PART 1
h2(doc, "Part 1 — The creator fills in the new-user wizard", chip="[ CREATOR ]", chip_color=GOLD)
para(doc, "Getting to the wizard:", bold=True, space_after=2)
step_table(doc, [
    (1, "An authorized leader signs in and opens Admin → Users tab."),
    (2, "If their role is Team Leader or higher, an “Add User” button appears at the top-right of the Users list."),
    (3, "Clicking “Add User” opens the New User wizard — a pop-up with three numbered steps."),
], "Creator — reaching the wizard", "8A6D1F")

para(doc, "Step 1 of 3 — Who is this person?", bold=True, space_after=2)
step_table(doc, [
    (4, "Type the person’s First name and Last name (both required)."),
    (5, "Enter an Email (required on this screen) and, optionally, a Phone number."),
    (6, "As the name is typed, the wizard auto-suggests a Username (e.g. “priscilla_aquila”). The creator can keep it or type their own; it must be at least 2 characters and not already taken (checked live)."),
    (7, "“Next” stays disabled until first name, last name, and email are filled and the username is valid and free."),
], "Creator — Step 1: identity", "8A6D1F")
figure(doc, "wizard-step1.png", "Step 1 collects name, email, optional phone, and an auto-suggested username.")

para(doc, "Step 2 of 3 — Where do they fit?", bold=True, space_after=2)
step_table(doc, [
    (8, "Pick a Role from a grid of buttons. The creator can only choose roles below their own level — a Developer sees all six roles; everyone else sees only the lower ones."),
    (9, "Pick who the new person Reports to (their parent in the org tree). It defaults to the creator (“you”); the list only includes active people ranked at or above the new person’s role."),
    (10, "Note: there is no church/area field and no group/team field here — the person’s church is worked out automatically from whoever they report to."),
], "Creator — Step 2: placement", "8A6D1F")
figure(doc, "wizard-step2.png", "Step 2 picks the role (limited to roles below the creator’s) and the “Reports to” parent.")

para(doc, "Step 3 of 3 — Review & confirm:", bold=True, space_after=2)
step_table(doc, [
    (11, "The wizard shows a read-only summary: name, username, email, phone, role, and who they report to."),
    (12, "An amber warning says a temporary password will be generated and shown only once — copy it before closing."),
    (13, "The creator clicks “Create account.”"),
], "Creator — Step 3: review", "8A6D1F")
figure(doc, "wizard-step3-review.png", "Step 3 reviews the values and warns that the temporary password is shown only once.")

# PART 2
h2(doc, "Part 2 — The system builds the account", chip="[ SYSTEM ]", chip_color=SLATE)
para(doc, "The moment “Create account” is clicked, the server runs a strict, security-ordered pipeline:")
step_table(doc, [
    (14, "Authenticate — confirm the creator is really signed in (from their session token, never anything typed in the form). No valid session → the request is refused (401)."),
    (15, "Authorize — re-check that this creator may create this role in this place (leaders only; you cannot create a role at or above your own; Branch Leaders and up can create anywhere; Team/Group Leaders only within their own group). Not allowed → refused (403)."),
    (16, "Validate the username — it must be 3–32 characters of letters, numbers, dot, dash, or underscore, and is stored in lowercase."),
    (17, "Check duplicates — a taken username → error (409); an already-used email → error (409). These run only AFTER the permission checks, on purpose, so an outsider can never fish for which usernames exist."),
    (18, "Check the “reports to” person — they must exist, be active, not be a plain Member, and outrank the new person."),
    (19, "Figure out the church/area — inherited by walking up the “reports to” chain to the nearest person who has a church (Overseers and Developers stay church-less)."),
    (20, "Save the account — marked Active, auto-tagged “teacher” for leader roles, and flagged “must change password” so the person is forced to set their own on first login."),
    (21, "Write an audit-log entry recording who created the account."),
    (22, "Generate a one-time temporary password (“Gc-…X9”), store it as the account’s real password, and hand it back to the wizard. The account is created (201)."),
], "System — server pipeline", "475569")

# PART 3
h2(doc, "Part 3 — The creator shares the credentials", chip="[ CREATOR ]", chip_color=GOLD)
step_table(doc, [
    (23, "The wizard flips to a green “Account created” screen showing the new Username and the temporary Password in plain text."),
    (24, "A “Copy both” button copies “Username: … / Password: …” to the clipboard."),
    (25, "The creator gives the username and temporary password to the new person directly. The app does not email them — delivery is manual/out-of-band."),
], "Creator — success & handoff", "8A6D1F")
figure(doc, "wizard-success.png", "The success screen shows the credentials once, with “Copy both.” The password here (Gc-vil8nf8sX9) is the live “Gc-…X9” format.")
callout(doc, "Shown once — there is no second chance",
        "The temporary password appears only on this success screen. If the creator closes the dialog without copying it, it cannot be recovered — the only fix is an admin “Reset password,” which issues a brand-new one.",
        CALL_AMBER, "C99A2E", AMBERTXT)

# PART 4
h2(doc, "Part 4 — The new user takes over the account", chip="[ NEW USER ]", chip_color=TEAL)
para(doc, "First sign-in:", bold=True, space_after=2)
step_table(doc, [
    (26, "The new person goes to the login page and signs in with their username (or email) and the temporary password they were given."),
    (27, "Login succeeds and briefly lands on the dashboard — but because the account is flagged “must change password,” the app immediately redirects to a “Set your password” screen and locks the rest of the app until they finish."),
], "New user — first sign-in", "0F766E")

para(doc, "Setting their own password (the forced change):", bold=True, space_after=2)
step_table(doc, [
    (28, "The “Set your password” screen greets them by first name and asks for a New password and a Confirm password — and nothing else."),
    (29, "The new password must be at least 8 characters and the two entries must match."),
    (30, "On submit, the server confirms they are changing their own password, saves it, and clears the “must change password” flag."),
    (31, "They are dropped back onto the dashboard — same session, now with full access."),
], "New user — set password", "0F766E")
figure(doc, "first-login.png", "The forced-change screen asks only for a new password (twice) — there is no “old/temporary password” field.")
callout(doc, "Common misconception — no “old password” box",
        "The forced-change screen does NOT ask for the temporary password again. The new user already proved it by logging in with it; this screen only asks for the new password twice. (There is also no “can’t reuse the temp password” rule.)",
        CALL_BLUE, "0F766E", TEAL)

para(doc, "Afterward:", bold=True, space_after=2)
step_table(doc, [
    (32, "The new password now signs them in normally (full dashboard access)."),
    (33, "The temporary password no longer works — any later sign-in attempt with it is rejected as “Invalid credentials.”"),
], "New user — aftermath", "0F766E")
figure(doc, "new-password-works.png", "With the new password set, the user signs in and reaches their dashboard (“Welcome back, Priscilla”).")
figure(doc, "old-password-rejected.png", "A later sign-in with the now-defunct temporary password is rejected — only the new password works.")

# ---- TEMP PASSWORD ----
h1(doc, "3. The temporary password, explained")
para(doc, "Every newly created account is born with a server-generated temporary password. It is real (login demands it), it is shown to the creator exactly once, and it is never stored by the browser afterward. Two different formats exist depending on how the account got its password:")
three_col(doc, ["Situation", "Format", "Example"],
    [
        ["New account (wizard)", "“Gc-” + 8 characters + “X9”", "Gc-vil8nf8sX9"],
        ["Contact → user conversion", "“Gc-” + 8 characters + “X9”", "Gc-3k9f1a02X9"],
        ["Admin “Reset password”", "Word + Word + 2 digits", "BrightRiver42"],
    ], widths=[2.4, 2.6, 1.8])
bullets(doc, [
    ("One-time reveal. ", "Shown once on the success (or reset) screen; unrecoverable afterward except by issuing a new one."),
    ("Forces a change. ", "The account is flagged “must change password,” so the temporary password only survives until the user sets their own."),
    ("Client vs. server rule. ", "The “Set your password” screen requires 8+ characters; the server itself accepts 6+. There is no “cannot reuse” check."),
])

# ---- ALTERNATIVE ENTRY POINTS ----
h1(doc, "4. Other ways an account is created")
h2(doc, "4.1  Converting a contact into a user")
para(doc, "On the Contacts page, opening a contact’s detail and clicking “Convert to user” creates a real login account in one step (a dedicated flow, not the wizard). It inherits the contact’s name, email, and phone; the operator picks only the role and “reports to” parent. The account is created with “must change password” set and its own temporary password, and the original contact is marked “converted” (kept for ~6 months). Converting an already-converted contact is blocked (no duplicate user).")
callout(doc, "Known gap",
        "The Contacts page currently confirms only the new @username after a conversion and does not surface the temporary password to the admin. In practice a converted user may need an admin “Reset password” to obtain a shareable credential.",
        CALL_AMBER, "C99A2E", AMBERTXT)
h2(doc, "4.2  Admin “Reset password”")
para(doc, "From the Users tab, the ⋮ menu on another user’s row (never your own) offers “Reset password.” It issues a fresh one-time temporary password (the “BrightRiver42” word-word-number format), shows it once with a copy button, and re-flags the account “must change password” so the user is sent through the same first-login change screen. Self password changes use a separate Settings flow instead.")

# ---- DATA WRITTEN ----
h1(doc, "5. What gets saved on a new account")
kv_table(doc, [
    ("Username", "Your chosen username, lowercased (3–32 chars: a–z, 0–9, dot, dash, underscore)"),
    ("Name / Email / Phone", "As entered (email is optional at the server; phone is optional)"),
    ("Role", "The role you selected"),
    ("Church / Area", "Inherited automatically from the “Reports to” chain (Overseers & Developers have none)"),
    ("Reports to (parent)", "The person you selected"),
    ("Tags", "“teacher” is added automatically for leader roles"),
    ("Status", "Active"),
    ("Must change password", "Yes — forces the first-login password change"),
    ("Temporary password", "Generated as “Gc-…X9”, stored as the real password, shown to you once (not saved on the record)"),
], label_w=2.1, col_hdr=["Field", "Value on a new account"])

# ---- CAVEATS ----
h1(doc, "6. Important caveats & gotchas")
bullets(doc, [
    ("Shown once. ", "If the temporary password isn’t copied from the success screen, it can’t be recovered — use “Reset password” to issue a new one."),
    ("Button vs. real permission. ", "Team and Group Leaders see “Add User” but can only actually create people inside their own group; the server enforces this on submit."),
    ("No “old password” box on first login. ", "The new user proves the temporary password by logging in; the “Set your password” screen only asks for the new password twice."),
    ("Temp password may be reused as the new one. ", "There’s no “can’t reuse” rule; the screen requires 8+ characters (the server accepts 6+)."),
    ("Delivery is manual. ", "The app doesn’t email credentials — the creator shares the username and temporary password directly."),
    ("First-login lock is client-side. ", "It’s enforced by the app screen, not the API, so expect a brief dashboard flash before the redirect."),
    ("Convert doesn’t reveal the temp password. ", "Converting a contact confirms the @username but not the password (see §4.1)."),
])

# ---- GLOSSARY ----
h1(doc, "7. Glossary")
kv_table(doc, [
    ("Role hierarchy", "Member < Team Leader < Group Leader < Branch Leader < Overseer < Developer."),
    ("Reports to (parent)", "The person directly above the new user in the org tree; determines their church."),
    ("Subtree", "Everyone beneath a leader in the org tree; sub-admins can only create within theirs."),
    ("Temporary password", "The one-time password an account is created with; the user must replace it on first login."),
    ("“Must change password” flag", "The marker that forces the first-login password screen and is cleared once the user sets their own."),
    ("Church / Area", "The location a user belongs to; inherited from their reporting chain."),
    ("Audit log", "The record of who did what; a “create” entry is written for every new account."),
], label_w=2.2)

# ---- SOURCE OF TRUTH ----
h1(doc, "8. Source of truth (code references)")
para(doc, "This article was written and verified against the source at build 5ce6222. If the behavior below changes, re-verify and bump the article version.", italic=True, color=MUTED, size=9.5)
three_col(doc, ["Area", "File", "Lines"],
    [
        ["Entry point (Users tab → Add User)", "src/components/admin/UsersTab.tsx", "440–445, 630–636"],
        ["Who can create (permissions)", "src/lib/utils/permissions.ts", "49–53, 182–207"],
        ["Create-user wizard (UI)", "src/components/users/CreateUserWizard.tsx", "49–421"],
        ["Create-user server handler", "src/mocks/handlers.ts", "2273–2453"],
        ["First-login page", "src/app/first-login/page.tsx", "37–124"],
        ["Change-password handler", "src/mocks/handlers.ts", "2821–2851"],
        ["Login + forced-change redirect", "use-auth.ts / (dashboard)/layout.tsx", "19–32 / 55–64"],
        ["Contact → user conversion", "ContactDetailDialog.tsx / handlers.ts", "668–749 / 2008–2125"],
        ["Admin reset password", "ResetPasswordDialog.tsx / handlers.ts", "– / 2778–2811"],
        ["Data model & contract tests", "types/user.ts, api/users.ts, *.itest.ts", "—"],
    ], widths=[2.5, 2.8, 1.5])

# ---- REVISION HISTORY ----
h1(doc, "9. Revision history")
three_col(doc, ["Version", "Date", "Change"],
    [
        ["v1.0", "2026-07-22", "Initial article. Verified against build 5ce6222; illustrated with live flow screenshots."],
    ], widths=[1.0, 1.3, 4.5])

doc.save(os.path.abspath(OUT))
print("SAVED:", os.path.abspath(OUT))
